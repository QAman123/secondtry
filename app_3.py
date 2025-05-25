import os
import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
import fitz  # PyMuPDF for PDF text extraction
import difflib
from fpdf import FPDF

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


def clean_text(text):
    replacements = {
        '\u2013': '-', '\u2014': '-',
        '\u2018': "'", '\u2019': "'",
        '\u201c': '"', '\u201d': '"',
        '\u2026': '...', '\u2022': '-',  # bullet point replaced by dash
        # add more replacements if needed
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text


def create_pdf(adapted_resume, cover_letter):
    adapted_resume = clean_text(adapted_resume)
    cover_letter = clean_text(cover_letter)

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Resume Page
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Adapted Resume", ln=True)
    pdf.set_font("Arial", '', 12)
    for line in adapted_resume.split('\n'):
        pdf.multi_cell(0, 8, line)

    # Cover Letter Page
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, "Cover Letter", ln=True)
    pdf.set_font("Arial", '', 12)
    for line in cover_letter.split('\n'):
        pdf.multi_cell(0, 8, line)

    # Generate PDF as bytes
    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    return BytesIO(pdf_bytes)


def create_docx(adapted_resume, cover_letter):
    doc = Document()
    doc.add_heading("Adapted Resume", level=1)
    for line in adapted_resume.split('\n'):
        doc.add_paragraph(line)

    doc.add_page_break()

    doc.add_heading("Cover Letter", level=1)
    for line in cover_letter.split('\n'):
        doc.add_paragraph(line)

    # Save to in-memory BytesIO
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream


def extract_text_from_pdf(uploaded_file):
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text


def extract_text_from_docx(file):
    try:
        doc = Document(file)
        fullText = []
        for para in doc.paragraphs:
            fullText.append(para.text)
        return "\n".join(fullText).strip()
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return None


def extract_resume_text(uploaded_file):
    if uploaded_file.type == "application/pdf":
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type in [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]:
        return extract_text_from_docx(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload PDF or DOCX.")
        return None


def generate_adapted_resume_and_cover(job_desc, resume_text, creative_instructions, language, temperature):
    system_prompt = (
        "You are an expert career coach and resume writer.\n"
        "Adapt the candidate's resume to better fit the job description.\n"
        "Then write a personalized cover letter highlighting relevant skills.\n"
        "Output both clearly separated with the headings:\n"
        "=== Adapted Resume ===\n"
        "and\n"
        "=== Cover Letter ==="
    )

    user_prompt = (
        f"Job Description:\n{job_desc}\n\nCandidate Resume:\n{resume_text}\n\n"
        f"{creative_instructions}\n\n"
        f"**Important:** Please provide the entire output in {language} language, regardless of the input language.\n\n"
        "Please provide:\n1) Adapted Resume\n2) Cover Letter\n\n"
        "Separate them exactly using the headings:\n"
        "=== Adapted Resume ===\n"
        "and\n"
        "=== Cover Letter ==="
    )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        # model="gpt-4o",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        max_tokens=3000,
        temperature=temperature,
    )
    return response.choices[0].message.content.strip()


def split_output(output_text):
    adapted_resume = ""
    cover_letter = ""

    if "=== Adapted Resume ===" in output_text and "=== Cover Letter ===" in output_text:
        parts = output_text.split("=== Cover Letter ===")
        adapted_resume = parts[0].replace("=== Adapted Resume ===", "").strip()
        cover_letter = parts[1].strip()
    else:
        adapted_resume = output_text

    return adapted_resume, cover_letter


PASSWORD = "two_cats"


def check_password():
    if "password_correct" not in st.session_state:
        st.session_state.password_correct = False

    if not st.session_state.password_correct:
        pwd = st.text_input("Enter password to access the app", type="password")
        if pwd:
            if pwd == PASSWORD:
                st.session_state.password_correct = True
                return True
            else:
                st.error("Incorrect password")
        return False
    else:
        return True


if check_password():
    st.title("Resume & Cover Letter Tailor")

    st.subheader("1. 💼 Paste the **Job Description** below:")
    job_desc = st.text_area("Job Description", height=200)

    st.subheader("2. 🎨 Choose Creative Enhancements")
    st.markdown(
        f"<span style='font-size: 0.85em; color: gray;'>You can pick multiple per category, but don't add too many to prevent inconsistency.</span>",
        unsafe_allow_html=True
    )

    # Style options
    style_options = [
        "Professional and polished",
        "Friendly and conversational",
        "Bold and confident",
        "Concise and direct",
        "Narrative storytelling",
        "Analytical and data-driven",
        "Warm and empathetic",
        "Visionary and inspiring",
        "Energetic and enthusiastic",
        "Calm and composed",
        "Sincere and humble",
        "Assertive and goal-oriented"
    ]

    # Flair options
    flair_options = [
        "Start with a short personal story",
        "Include subtle humor",
        "Include humor",
        "Incorporate motivational language",
        "Highlight entrepreneurial mindset",
        "Showcase emotional intelligence",
        "Emphasize results over responsibilities",
        "Add a visionary or future-forward tone",
        "Use metaphors or analogies",
        "Tie in a relevant quote or motto",
        "Include a call to action at the end",
        "Weave in a unique career insight",
        "Celebrate cultural intelligence or diversity",
        "Connect past roles to broader global trends",
        "Mention meaningful personal passions (if relevant)"
    ]

    # Company type options
    company_type_options = [
        "Tailor for a startup company",
        "Tailor for a corporate enterprise",
        "Appeal to a creative industry",
        "Fit for a non-profit/mission-driven org",
        "Align with a fast-scaling tech company",
        "Fit for a government/public sector role",
        "Appeal to a global/multinational organization",
        "Tailor for a consultancy or agency environment",
        "Fit for a remote-first or flexible workplace",
        "Appeal to a data-driven and analytical culture"
    ]

    # Weights options dropdown
    weight_labels = ["Low", "Medium", "High"]
    weight_values = {"Low": 1, "Medium": 3, "High": 5}

    # Create three columns for categories
    col1, col2, col3 = st.columns(3)

    # Dictionary to store selected enhancements and their weights
    enhancement_weights = {}

    with col1:
        st.markdown("**Writing Style**")
        selected_style = st.multiselect(
            "Choose writing style:",
            style_options,
            key="style_select"
        )
        for style in selected_style:
            weight = st.selectbox(
                f"Weight for '{style}'",
                weight_labels,
                index=1,
                key=f"weight_style_{style}"
            )
            enhancement_weights[style] = weight_values[weight]

    with col2:
        st.markdown("**Creative Flair**")
        selected_flair = st.multiselect(
            "Add creative elements:",
            flair_options,
            key="flair_select"
        )
        for flair in selected_flair:
            weight = st.selectbox(
                f"Weight for '{flair}'",
                weight_labels,
                index=1,
                key=f"weight_flair_{flair}"
            )
            enhancement_weights[flair] = weight_values[weight]

    with col3:
        st.markdown("**Company Type**")
        selected_company = st.multiselect(
            "Target company type:",
            company_type_options,
            key="company_select"
        )
        for comp in selected_company:
            weight = st.selectbox(
                f"Weight for '{comp}'",
                weight_labels,
                index=1,
                key=f"weight_company_{comp}"
            )
            enhancement_weights[comp] = weight_values[weight]

    # Build creative instructions string with weights
    if enhancement_weights:
        creative_instructions = "Creative Enhancements to apply (with importance weights):\n" + "\n".join(
            f"- {enhancement} (importance: {weight})" for enhancement, weight in enhancement_weights.items()
        )
    else:
        creative_instructions = ""

    st.subheader("3. 🎯 AI Creativity Level")
    creativity_options = {
        "Low": 0.3,
        "Medium": 0.7,
        "High": 1.0,
        "Very High": 1.3
    }

    creativity_level = st.selectbox(
        "Select creativity level for AI responses:",
        options=list(creativity_options.keys()),
        index=1,  # Default to "Medium"
        help="Higher creativity may produce more varied and creative outputs, but lower creativity ensures more consistent and focused results."
    )

    temperature = creativity_options[creativity_level]

    st.subheader("4. 💬 What language do you want to work in?")
    language = st.selectbox(
        "Select output language",
        options=["English", "Dutch", "Spanish", "French", "German", "Chinese", "Japanese", "Russian"],
        index=0,
    )

    st.subheader("5. 📎 Upload your Resume (PDF or DOCX)")
    uploaded_file = st.file_uploader("", type=["pdf", "docx"])

    if uploaded_file:
        resume_text = extract_resume_text(uploaded_file)
        if resume_text:
            st.markdown("Extracted Resume Text Preview")
            st.text_area("Resume Text", resume_text[:8000], height=100)

            if st.button("🔔 6. Generate Adapted Resume & Cover Letter"):
                with st.spinner("Generating..."):
                    output = generate_adapted_resume_and_cover(job_desc, resume_text, creative_instructions, language,
                                                               temperature)
                    adapted_resume, cover_letter = split_output(output)

                st.session_state["adapted_resume"] = adapted_resume
                st.session_state["cover_letter"] = cover_letter

    if "adapted_resume" in st.session_state and "cover_letter" in st.session_state:
        st.subheader("******* 👍👍 Results 👍👍 *******")
        st.subheader("Suggested Resume (check carefully for accuracy!)")
        st.text_area("", st.session_state["adapted_resume"], height=600)

        st.subheader("Cover Letter (check carefully for accuracy!)")
        st.text_area("", st.session_state["cover_letter"], height=600)

        docx_file = create_docx(st.session_state["adapted_resume"], st.session_state["cover_letter"])
        pdf_file = create_pdf(st.session_state["adapted_resume"], st.session_state["cover_letter"])

        st.download_button(
            label="Download Adapted Resume & Cover Letter as DOCX",
            data=docx_file,
            file_name="adapted_resume_and_cover_letter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    else:
        st.info("Please upload a resume and enter job description to start.")